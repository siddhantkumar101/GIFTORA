from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

def shade_row(row, c="FFF7ED"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), c); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'CCCCCC')
    pBdr.append(b); pPr.append(pBdr)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(249,115,22)
def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(30,30,30)
def para(t, bold=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor(30,30,30)
def bullet(t, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if prefix:
        r1 = p.add_run(prefix+": "); r1.bold = True; r1.font.size = Pt(11)
        p.add_run(t).font.size = Pt(11)
    else:
        p.add_run(t).font.size = Pt(11)
def code_block(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
def orange_header(table, headers):
    shade_row(table.rows[0], "F97316")
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        r = table.rows[0].cells[i].paragraphs[0].runs[0]
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255)

# ═══════════════════ COVER ═══════════════════
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("GIFTORA STUDIO"); tr.bold = True; tr.font.size = Pt(32)
tr.font.color.rgb = RGBColor(249,115,22)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("Milestone 2 Delivery Report"); sr.bold = True; sr.font.size = Pt(18)
doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
subr = sub.add_run("E-commerce System + Admin Panel + Deployment")
subr.italic = True; subr.font.size = Pt(13); subr.font.color.rgb = RGBColor(100,100,100)
doc.add_paragraph(); doc.add_paragraph()

info = doc.add_table(rows=6, cols=2); info.style = 'Table Grid'
info.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (l,v) in enumerate([
    ("Project",   "Giftora Studio — Custom Gifting Platform"),
    ("Milestone", "Milestone 2: E-commerce System + Admin Panel + Deployment"),
    ("Amount",    "$600"),
    ("Live URL",  "https://giftora-six.vercel.app"),
    ("GitHub",    "https://github.com/siddhantkumar101/GIFTORA"),
    ("Date",      datetime.date.today().strftime("%B %d, %Y")),
]):
    info.rows[i].cells[0].text = l; info.rows[i].cells[1].text = v
    info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(info.rows[i])
doc.add_page_break()

# ═══════════════════ 1. EXECUTIVE SUMMARY ═══════════════════
h1("1. Executive Summary"); add_hr()
para("Milestone 2 delivers the complete e-commerce engine — including dynamic product catalog "
     "management, a full cart and checkout flow with pricing logic, a secure payment confirmation "
     "system, end-to-end order tracking with a 6-stage fulfillment pipeline, and a comprehensive "
     "Admin Dashboard for managing products, pricing, inventory, and order fulfillment.")
doc.add_paragraph()
para("The platform is live at https://giftora-six.vercel.app with MongoDB Atlas persistence, "
     "automatic CI/CD via GitHub + Vercel, and full role-based access control.")
doc.add_paragraph()

# ═══════════════════ 2. SCOPE ═══════════════════
h1("2. Scope of Work Delivered"); add_hr()
for title, desc in [
    ("Product Catalog Management",
     "Sellers can create, update, and delete products from the Admin Dashboard. Products support "
     "name, category (12 categories), price, image (file upload, clipboard paste, or URL), description, "
     "material, and lead time. Changes are reflected instantly across the platform."),
    ("Cart and Checkout System",
     "A full shopping cart allows customers to add customized gifts with all metadata preserved, "
     "update quantities inline, remove items, and proceed to checkout. The checkout collects delivery "
     "address, validates all required fields, and submits orders to the backend API."),
    ("Payment Gateway Integration",
     "A simulated secure payment system supports three methods: Giftora Secure Demo Pay, Card test "
     "payment, and UPI test payment. Automatic pricing logic: Rs. 79 delivery waived above Rs. 999, "
     "Rs. 150 discount applied above Rs. 1,499. Each order generates a unique transaction ID."),
    ("Order Tracking System",
     "After checkout, orders receive a unique order number (format: GFT-XXXXXX-XXX) and are persisted "
     "to MongoDB with full customer, items, totals, payment, and tracking data. Customers view all "
     "orders filtered by their email on the Orders page."),
    ("Admin Dashboard",
     "A seller-only dashboard provides 4 real-time metric cards (Orders, Revenue, Catalog size, Status), "
     "a Live Inventory list with edit/delete controls, and a New Catalog Item form supporting image "
     "upload via file picker, clipboard paste (Ctrl+V), or URL input."),
    ("Product and Order Management",
     "Sellers view all customer orders in an Order Queue sidebar and update status through a 6-stage "
     "pipeline: Design approval → Printing → Quality check → Packed → Shipped → Delivered. Each "
     "status change is recorded in a tracking array with timestamp and note."),
    ("Testing and Final Deployment",
     "The application is deployed on Vercel (frontend) with MongoDB Atlas (database). Backend includes "
     "Jest + Supertest testing infrastructure with cross-env for cross-platform ENV support. The app "
     "handles all error states gracefully."),
]:
    bullet(desc, prefix=title); doc.add_paragraph()

# ═══════════════════ 3. ENVIRONMENT CONFIGURATION ═══════════════════
h1("3. Environment & Configuration"); add_hr()

h2("3.1 Server Environment Variables (.env)")
para("The server requires the following environment variables defined in server/.env:")
doc.add_paragraph()
env_table = doc.add_table(rows=5, cols=3); env_table.style = 'Table Grid'
orange_header(env_table, ["Variable", "Example Value", "Description"])
for i, (var, val, desc) in enumerate([
    ("PORT",         "5000",                                    "Port the Express server listens on"),
    ("CLIENT_ORIGIN","http://localhost:5173",                    "Allowed CORS origin for the frontend"),
    ("MONGODB_URI",  "mongodb+srv://user:pass@cluster.mongodb.net/giftora", "MongoDB Atlas connection string"),
    ("JWT_SECRET",   "giftora_studio_secret_key_2026",          "Secret key for signing JWT tokens (7-day expiry)"),
]):
    r = env_table.rows[i+1]
    r.cells[0].text = var; r.cells[1].text = val; r.cells[2].text = desc
    r.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

h2("3.2 Cookie Security Configuration")
cookie_t = doc.add_table(rows=6, cols=2); cookie_t.style = 'Table Grid'
orange_header(cookie_t, ["Property", "Value"])
for i, (p2,v) in enumerate([
    ("httpOnly", "true — prevents client-side JavaScript access (XSS protection)"),
    ("secure",   "true — cookie transmitted only over HTTPS"),
    ("sameSite", "none — allows cross-origin cookies for Vercel deployment"),
    ("path",     "/ — cookie available on all routes"),
    ("maxAge",   "7 days (604,800,000 ms)"),
]):
    r = cookie_t.rows[i+1]
    r.cells[0].text = p2; r.cells[1].text = v
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

h2("3.3 CORS Policy")
para("The server accepts credentials from any origin to support the decoupled Vercel frontend + API model:")
code_block("""app.use(cors({
  origin: (origin, callback) => callback(null, true),
  credentials: true
}));""")
doc.add_paragraph()

h2("3.4 Vite Dev Proxy")
para("In development, the Vite config proxies /api requests to the local Express server:")
code_block("""server: {
  port: 5173,
  proxy: { "/api": "http://localhost:5000" }
}""")
doc.add_paragraph()

# ═══════════════════ 4. CART & CHECKOUT ═══════════════════
h1("4. Cart & Checkout System — Implementation Details"); add_hr()

h2("4.1 Cart Features")
for b in [
    "Items from Studio carry full customization metadata (text, font, color, placement, image, quantity)",
    "Cart persisted to localStorage — survives page refreshes and browser restarts",
    "Quantity updated inline per item with live subtotal recalculation",
    "Individual item removal with one click",
    "Empty cart state with clear prompt to return to Studio",
    "Cart item count badge visible in both desktop header and mobile bottom nav",
]:
    bullet(b)
doc.add_paragraph()

h2("4.2 Pricing Logic")
pl = doc.add_table(rows=4, cols=2); pl.style = 'Table Grid'
orange_header(pl, ["Rule", "Value"])
for i, (rule,val) in enumerate([
    ("Delivery fee (standard)",   "Rs. 79"),
    ("Free delivery threshold",   "Orders above Rs. 999"),
    ("Bulk discount",             "Rs. 150 off on orders above Rs. 1,499"),
]):
    r = pl.rows[i+1]; r.cells[0].text = rule; r.cells[1].text = val
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

h2("4.3 Checkout Flow")
para("1. User must be logged in as Consumer → 2. Cart must not be empty → "
     "3. Name + Email required → 4. Address line1 + City required → "
     "5. Payment method selected → 6. Order submitted via POST /api/orders → "
     "7. Order number generated → 8. Cart cleared → 9. Redirected to Orders page")
doc.add_paragraph()

# ═══════════════════ 5. ORDER SYSTEM ═══════════════════
h1("5. Order Tracking System"); add_hr()

h2("5.1 Order Schema (MongoDB — from actual code)")
code_block("""\
{
  orderNumber:  String  (unique, indexed)          — Auto-generated: "GFT-{timestamp}-{random}"
  customer: {
    name:       String  (required),
    email:      String  (required, indexed),
    phone:      String  (default: "")
  },
  address: {
    line1:      String  (required),
    city:       String  (required),
    state:      String  (required),
    pincode:    String  (required)
  },
  items: [{
    productSlug:    String,
    name:           String,
    price:          Number,
    quantity:       Number,
    customization:  Mixed     — { text, textColor, productColor, placement, font, size, imageName, hasImage }
  }],
  totals: {
    subtotal:   Number,
    delivery:   Number,
    discount:   Number,
    grandTotal: Number
  },
  payment: {
    method:        String,
    status:        String  (default: "Paid in demo mode"),
    transactionId: String  — Auto-generated: "TXN-{timestamp}"
  },
  status:    String  (default: "Design approval"),
  tracking:  [{
    label:   String,
    at:      Date,
    note:    String
  }],
  timestamps: createdAt, updatedAt
}""")
doc.add_paragraph()

h2("5.2 Fulfillment Status Pipeline")
st = doc.add_table(rows=7, cols=2); st.style = 'Table Grid'
orange_header(st, ["Stage", "Description"])
for i, (stage,desc) in enumerate([
    ("Design approval", "Order received, design reviewed and confirmed"),
    ("Printing",        "Product being printed / engraved / crafted"),
    ("Quality check",   "Product inspected for defects before packing"),
    ("Packed",          "Securely packed with gift-ready packaging"),
    ("Shipped",         "Dispatched via courier, tracking active"),
    ("Delivered",       "Order received by customer — order complete"),
]):
    r = st.rows[i+1]; r.cells[0].text = stage; r.cells[1].text = desc
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 6. ADMIN DASHBOARD ═══════════════════
h1("6. Admin Dashboard"); add_hr()

h2("6.1 Access Control")
para("The Admin Dashboard is exclusively accessible to 'seller' role users. Unauthorized access "
     "to /admin is blocked by React Router with automatic redirect to homepage. The admin login "
     "uses the same JWT system with role validation to prevent privilege escalation.")
doc.add_paragraph()

h2("6.2 Analytics Metrics (4 Real-Time Cards)")
mt = doc.add_table(rows=5, cols=2); mt.style = 'Table Grid'
orange_header(mt, ["Metric", "Data Source"])
for i, (m,s2) in enumerate([
    ("Orders",  "Live count of all orders in MongoDB"),
    ("Revenue", "Sum of grandTotal across all orders (formatted as Rs.)"),
    ("Catalog", "Total number of active products in inventory"),
    ("Status",  "Platform health indicator — shows 'Online'"),
]):
    r = mt.rows[i+1]; r.cells[0].text = m; r.cells[1].text = s2
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

h2("6.3 Product Management Features")
for b in [
    "Create products via full-featured form: name, category (12 options), price, image, description, material, lead time",
    "Image upload supports 3 methods: local file picker, clipboard paste (Ctrl+V), and remote URL input",
    "Edit any product — form pre-fills current values for easy inline updating",
    "Delete products permanently with confirmation dialog",
    "All changes reflected in real time across the storefront without page reload",
    "Live Inventory list with image thumbnails, name, category, and price for quick scanning",
]:
    bullet(b)
doc.add_paragraph()

h2("6.4 Order Queue Management")
for b in [
    "All customer orders displayed in sidebar Order Queue, sorted by creation time",
    "Each order card shows: order number, customer name, item count, total value",
    "Status updated via dropdown — changes saved instantly to MongoDB with tracking entry",
    "6-stage pipeline provides clear workflow visibility for fulfillment teams",
    "Empty state handled gracefully when no orders are pending",
]:
    bullet(b)
doc.add_paragraph()

# ═══════════════════ 7. API ENDPOINTS ═══════════════════
h1("7. API Endpoints — Milestone 2"); add_hr()
api = doc.add_table(rows=8, cols=3); api.style = 'Table Grid'
orange_header(api, ["Method + Route", "Description", "Auth"])
for i, (m,d,a) in enumerate([
    ("GET /api/orders",                       "Fetch all (seller) or by email (consumer)",    "Yes"),
    ("POST /api/orders",                      "Place order with items, totals, address",       "Consumer"),
    ("PATCH /api/orders/:orderNumber/status",  "Update order fulfillment status",              "Seller"),
    ("POST /api/products",                    "Create new product",                            "Seller"),
    ("PATCH /api/products/:slug",             "Update product fields",                         "Seller"),
    ("DELETE /api/products/:slug",            "Delete product permanently",                    "Seller"),
    ("GET /api/health",                       "Server + database health check",                "No"),
]):
    r = api.rows[i+1]
    r.cells[0].text = m; r.cells[1].text = d; r.cells[2].text = a
    r.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 8. TESTING ═══════════════════
h1("8. Testing"); add_hr()

h2("8.1 Backend Testing Setup")
para("The backend includes Jest + Supertest testing configured for ES Module support. Tests run "
     "in NODE_ENV=test which bypasses database connection and uses in-memory store.")
code_block(""""test": "cross-env NODE_ENV=test node --experimental-vm-modules node_modules/jest/bin/jest.js"

Dependencies: jest ^30.4.2, supertest ^7.2.2, cross-env ^10.1.0""")
doc.add_paragraph()

h2("8.2 Manual Testing Scenarios Verified")
for b in [
    "Place order as logged-in consumer → order saved with correct totals and tracking entry",
    "Place order without address → validation error, order not saved",
    "Place order as seller → blocked (role mismatch)",
    "Add product as seller → appears in catalog immediately",
    "Update product price as seller → reflected in Studio and Home pages",
    "Delete product → removed from all listings",
    "Update order status → status saved, tracking entry added with timestamp",
    "Access /admin as consumer → redirected to homepage",
    "Cart persistence → add items, close browser, reopen → cart intact",
    "Free delivery threshold (order > Rs. 999) → delivery = Rs. 0",
    "Bulk discount (order > Rs. 1,499) → Rs. 150 off applied",
    "Server cold start → fallback data shown instantly, real data loads silently",
]:
    bullet(b)
doc.add_paragraph()

# ═══════════════════ 9. DEPLOYMENT ═══════════════════
h1("9. Deployment Architecture"); add_hr()
dt = doc.add_table(rows=7, cols=2); dt.style = 'Table Grid'
orange_header(dt, ["Component", "Details"])
for i, (c2,d) in enumerate([
    ("Live URL",             "https://giftora-six.vercel.app"),
    ("Frontend Hosting",     "Vercel — auto-deploys from GitHub main branch"),
    ("Build Tool",           "Vite 8 — optimized production build with tree-shaking"),
    ("Backend / API",        "Vercel Serverless via vercel.json routing"),
    ("Database",             "MongoDB Atlas — cloud cluster, free tier (M0)"),
    ("CI/CD",                "GitHub → Vercel webhook — zero-downtime deployments"),
]):
    r = dt.rows[i+1]; r.cells[0].text = c2; r.cells[1].text = d
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 10. CHECKLIST ═══════════════════
h1("10. Milestone 2 Delivery Checklist"); add_hr()
chk = doc.add_table(rows=8, cols=2); chk.style = 'Table Grid'
orange_header(chk, ["Feature", "Status"])
for i, (f,s2) in enumerate([
    ("Product catalog management",   "✅ Complete — Admin CRUD for products"),
    ("Cart and checkout system",     "✅ Complete — Full cart + validated checkout"),
    ("Payment gateway integration",  "✅ Complete — 3 methods, pricing logic, transaction IDs"),
    ("Order tracking system",        "✅ Complete — MongoDB persistence, customer view, tracking"),
    ("Admin dashboard",              "✅ Complete — Metrics, inventory, order queue"),
    ("Product and order management", "✅ Complete — CRUD products, 6-stage order pipeline"),
    ("Testing and final deployment", "✅ Complete — Jest/Supertest, live on Vercel + Atlas"),
]):
    r = chk.rows[i+1]; r.cells[0].text = f; r.cells[1].text = s2
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)

doc.add_paragraph(); doc.add_paragraph()
cl = doc.add_paragraph(); cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cl.add_run("All deliverables for Milestone 2 have been completed, tested, and deployed. "
                "The Giftora platform is fully operational end-to-end.")
cr.italic = True; cr.font.size = Pt(11); cr.font.color.rgb = RGBColor(80,80,80)

doc.save(r"C:\Users\siddh\Desktop\Giftora_Milestone2_Report_v2.docx")
print("Milestone 2 saved!")
