from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def shade_row(row, c="FFF7ED"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), c)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), 'CCCCCC')
    pBdr.append(b); pPr.append(pBdr)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(249,115,22)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(30,30,30)

def para(t, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(30,30,30)
    return p

def bullet(t, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if prefix:
        r1 = p.add_run(prefix + ": "); r1.bold = True; r1.font.size = Pt(11)
        p.add_run(t).font.size = Pt(11)
    else:
        p.add_run(t).font.size = Pt(11)

def code_block(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.name = "Courier New"; r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)

def orange_header(table, headers):
    shade_row(table.rows[0], "F97316")
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        r = table.rows[0].cells[i].paragraphs[0].runs[0]
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255)

# ═══════════════════ COVER PAGE ═══════════════════
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("GIFTORA STUDIO"); tr.bold = True; tr.font.size = Pt(32)
tr.font.color.rgb = RGBColor(249,115,22)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("Milestone 1 Delivery Report"); sr.bold = True; sr.font.size = Pt(18)

doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
subr = sub.add_run("UI/UX Design + Product Customization System")
subr.italic = True; subr.font.size = Pt(13); subr.font.color.rgb = RGBColor(100,100,100)

doc.add_paragraph(); doc.add_paragraph()
info = doc.add_table(rows=6, cols=2); info.style = 'Table Grid'
info.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (l,v) in enumerate([
    ("Project", "Giftora Studio — Custom Gifting Platform"),
    ("Milestone", "Milestone 1: UI/UX Design + Product Customization System"),
    ("Amount", "$600"),
    ("Live URL", "https://giftora-six.vercel.app"),
    ("GitHub", "https://github.com/siddhantkumar101/GIFTORA"),
    ("Date", datetime.date.today().strftime("%B %d, %Y")),
]):
    info.rows[i].cells[0].text = l; info.rows[i].cells[1].text = v
    info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(info.rows[i])

doc.add_page_break()

# ═══════════════════ 1. EXECUTIVE SUMMARY ═══════════════════
h1("1. Executive Summary"); add_hr()
para("Milestone 1 of the Giftora Studio project has been successfully completed. This milestone "
     "covers the complete design and implementation of the platform's core foundation — including "
     "a fully responsive UI with mobile-first design, a real-time product customization engine, "
     "user authentication system with JWT + secure cookies, image upload functionality, live product "
     "preview, and a fully configured MERN stack backend with MongoDB Atlas persistence.")
doc.add_paragraph()
para("The platform is live and publicly accessible at https://giftora-six.vercel.app. All features "
     "outlined in the Milestone 1 scope have been implemented, tested, and deployed.")
doc.add_paragraph()

# ═══════════════════ 2. SCOPE OF WORK ═══════════════════
h1("2. Scope of Work Delivered"); add_hr()
for title, desc in [
    ("Homepage and Product Pages",
     "A fully responsive homepage with hero section, dynamic product categories derived from the live catalog, bestseller grid, testimonials, and a call-to-action section. Mobile-first design with a dedicated MobileHome component for touch-optimized navigation."),
    ("User Authentication System",
     "Complete JWT-based authentication with login, registration, and logout flows for both Consumer and Admin (Seller) roles. Sessions are persisted via secure HTTP-only cookies with SameSite=None and Secure flag for cross-origin cookie support. Passwords are hashed with bcryptjs (10 salt rounds). Tokens expire after 7 days."),
    ("Product Customization Module",
     "An interactive Studio page allowing users to browse the product catalog, select a product, and apply full customizations including custom text, text color, product color, placement, font style, and quantity."),
    ("Image Upload Functionality",
     "Users can upload a personal image (logo, photo, artwork) directly in the browser. The uploaded image is processed client-side using the FileReader API and immediately reflected in the live preview."),
    ("Custom Text Editing",
     "A rich text customization panel allows users to enter up to 72 characters of custom text, select font styles (Sans, Serif, Mono, Script) and text color via a color picker."),
    ("Live Product Preview System",
     "A real-time PreviewPanel renders the selected product alongside all applied customizations so the customer sees exactly what their final product will look like before ordering."),
    ("Project Setup and Database Configuration",
     "Full MERN stack architecture configured with MongoDB Atlas for persistence. The server gracefully falls back to in-memory seeded data if the database is unavailable, ensuring zero downtime."),
]:
    bullet(desc, prefix=title); doc.add_paragraph()

# ═══════════════════ 3. TECH STACK ═══════════════════
h1("3. Technology Stack"); add_hr()
st = doc.add_table(rows=8, cols=3); st.style = 'Table Grid'
orange_header(st, ["Layer", "Technology", "Purpose"])
for i, (l,t2,p2) in enumerate([
    ("Frontend",   "React 18 + Vite 8",         "Component-based SPA with fast HMR and optimized builds"),
    ("Styling",    "Tailwind CSS 3",             "Utility-first responsive design system"),
    ("Routing",    "React Router DOM v7",        "Client-side navigation and protected routes"),
    ("Icons",      "Lucide React",               "Consistent SVG icon library"),
    ("Backend",    "Node.js + Express 4",        "RESTful API server with middleware support"),
    ("Database",   "MongoDB Atlas + Mongoose 8", "Cloud-hosted document database with schema validation"),
    ("Auth",       "JWT + bcryptjs + Cookies",   "Secure stateless auth with encrypted passwords"),
]):
    r = st.rows[i+1]
    r.cells[0].text = l; r.cells[1].text = t2; r.cells[2].text = p2
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 4. ENVIRONMENT CONFIGURATION ═══════════════════
h1("4. Environment & Configuration"); add_hr()

h2("4.1 Server Environment Variables (.env)")
para("The server requires the following environment variables, defined in server/.env:")
doc.add_paragraph()
env_table = doc.add_table(rows=5, cols=3); env_table.style = 'Table Grid'
orange_header(env_table, ["Variable", "Example Value", "Description"])
for i, (var, val, desc) in enumerate([
    ("PORT",         "5000",                                    "Port the Express server listens on"),
    ("CLIENT_ORIGIN","http://localhost:5173",                    "Allowed CORS origin for the frontend"),
    ("MONGODB_URI",  "mongodb+srv://user:pass@cluster.mongodb.net/giftora", "MongoDB Atlas connection string (required for persistence)"),
    ("JWT_SECRET",   "giftora_studio_secret_key_2026",          "Secret key used to sign and verify JWT tokens"),
]):
    r = env_table.rows[i+1]
    r.cells[0].text = var; r.cells[1].text = val; r.cells[2].text = desc
    r.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

h2("4.2 Client Configuration (vite.config.js)")
para("The Vite configuration includes a development proxy to forward API requests to the backend:")
code_block("""export default defineConfig({
  plugins: [react()],
  server: {
    port: 5173,
    proxy: {
      "/api": "http://localhost:5000"
    }
  }
});""")
doc.add_paragraph()
para("In production, the frontend uses the VITE_API_URL environment variable (or defaults to /api) "
     "configured in the centralized api.js utility to connect to the deployed backend.")
doc.add_paragraph()

h2("4.3 Cookie Configuration")
para("Authentication cookies are configured with security best practices for cross-origin deployment:")
cookie_table = doc.add_table(rows=6, cols=2); cookie_table.style = 'Table Grid'
orange_header(cookie_table, ["Property", "Value"])
for i, (prop, val) in enumerate([
    ("httpOnly", "true — prevents JavaScript access (XSS protection)"),
    ("secure",   "true — cookie only sent over HTTPS"),
    ("sameSite", "none — allows cross-origin cookies (required for Vercel → separate API)"),
    ("path",     "/ — cookie available across all routes"),
    ("maxAge",   "7 days (604,800,000 ms) — auto-expiry for sessions"),
]):
    r = cookie_table.rows[i+1]
    r.cells[0].text = prop; r.cells[1].text = val
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

h2("4.4 CORS Configuration")
para("The server accepts requests from any origin with credentials enabled, to support the "
     "cross-site Vercel deployment model:")
code_block("""app.use(cors({
  origin: (origin, callback) => callback(null, true),
  credentials: true
}));""")
doc.add_paragraph()

h2("4.5 Database Fallback Strategy")
para("If MONGODB_URI is not set or the database connection fails, the server automatically "
     "switches to an in-memory store pre-populated with 15 seeded products. This ensures the "
     "application works end-to-end in demo mode without any external dependencies.")
doc.add_paragraph()

# ═══════════════════ 5. FILE STRUCTURE ═══════════════════
h1("5. Project File Structure"); add_hr()
code_block("""\
giftora/
├── client/                        # React + Vite Frontend
│   ├── src/
│   │   ├── components/
│   │   │   ├── CategoryCard.jsx   # Shop-by-category card
│   │   │   ├── CustomizerPanel.jsx # Product customization form
│   │   │   ├── Footer.jsx         # Site footer
│   │   │   ├── LoginPanel.jsx     # Auth modal (login/register)
│   │   │   ├── MobileHome.jsx     # Mobile-optimized homepage
│   │   │   ├── PreviewPanel.jsx   # Live product preview
│   │   │   ├── ProductCard.jsx    # Product grid card
│   │   │   └── Testimonials.jsx   # Customer reviews section
│   │   ├── pages/
│   │   │   ├── AdminView.jsx      # Seller dashboard
│   │   │   ├── CartView.jsx       # Shopping cart + checkout
│   │   │   ├── HomeView.jsx       # Main homepage
│   │   │   ├── OrdersView.jsx     # Customer order history
│   │   │   └── StudioView.jsx     # Product customization studio
│   │   ├── utils/
│   │   │   ├── api.js             # Centralized fetch wrapper with credentials
│   │   │   ├── constants.js       # Fallback data, placements, fonts, defaults
│   │   │   └── helpers.js         # Money formatter, image optimizer
│   │   ├── App.jsx                # Root component + routing + state management
│   │   ├── index.css              # Global Tailwind styles + custom tokens
│   │   └── main.jsx               # React DOM entry point
│   ├── index.html
│   ├── package.json
│   ├── tailwind.config.js
│   └── vite.config.js             # Dev proxy: /api → localhost:5000
│
├── server/                        # Node.js + Express Backend
│   ├── config/
│   │   └── db.js                  # MongoDB connection + in-memory fallback
│   ├── controllers/
│   │   ├── authController.js      # Register, login, logout, getMe (JWT)
│   │   ├── orderController.js     # Order CRUD + status pipeline
│   │   └── productController.js   # Product CRUD + lite/full modes
│   ├── data/
│   │   └── seedData.js            # 15 pre-seeded products + status flow
│   ├── models/
│   │   ├── Order.js               # Order schema (customer, items, totals, tracking)
│   │   ├── Product.js             # Product schema (slug, colors, areas, etc.)
│   │   └── User.js                # User schema (bcrypt password, roles, addresses)
│   ├── routes/
│   │   ├── authRoutes.js          # POST register/login/logout, GET /me
│   │   ├── orderRoutes.js         # GET/POST orders, PATCH status
│   │   └── productRoutes.js       # GET/POST products, PATCH/DELETE by slug
│   ├── seed.js                    # Standalone DB seeder script
│   ├── server.js                  # Express app entry + middleware + static serve
│   ├── package.json
│   ├── .env                       # Environment variables (not committed)
│   └── .env.example               # Template for required env vars
│
├── vercel.json                    # Vercel deployment: framework=vite, root=client
├── .gitignore                     # node_modules, .env, dist excluded
└── README.md""")
doc.add_paragraph()

# ═══════════════════ 6. API ENDPOINTS ═══════════════════
h1("6. API Endpoints Implemented"); add_hr()
api = doc.add_table(rows=11, cols=3); api.style = 'Table Grid'
orange_header(api, ["Method + Route", "Description", "Auth"])
for i, (m,d,a) in enumerate([
    ("GET /api/health",             "Server health check + DB status",               "No"),
    ("GET /api/products",           "Fetch all active products (full payload)",       "No"),
    ("GET /api/products?lite=true", "Fetch products with reduced payload (fast)",     "No"),
    ("POST /api/products",          "Create new product in catalog",                  "Seller"),
    ("PATCH /api/products/:slug",   "Update product fields by slug",                 "Seller"),
    ("DELETE /api/products/:slug",  "Remove product permanently",                    "Seller"),
    ("POST /api/auth/register",     "Register new user (consumer or seller)",        "No"),
    ("POST /api/auth/login",        "Login and receive session cookie",              "No"),
    ("POST /api/auth/logout",       "Clear session cookie",                          "Yes"),
    ("GET /api/auth/me",            "Verify session and return current user",        "Yes"),
]):
    r = api.rows[i+1]
    r.cells[0].text = m; r.cells[1].text = d; r.cells[2].text = a
    r.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 7. DATABASE SCHEMAS ═══════════════════
h1("7. Database Schemas (MongoDB)"); add_hr()

h2("7.1 Product Schema")
code_block("""\
{
  slug:               String  (unique, indexed)      — URL-safe product identifier
  name:               String  (required)             — Display name
  category:           String  (required)             — e.g. "Mugs", "Corporate", "Luxury"
  price:              Number  (required)             — Price in INR
  compareAt:          Number  (default: 0)           — Strikethrough / MRP price
  image:              String  (required)             — Product image URL (Unsplash)
  description:        String  (required)             — Product description
  leadTime:           String  (required)             — e.g. "3 days"
  rating:             Number  (default: 5)           — Star rating (out of 5)
  orders:             Number  (default: 0)           — Number of orders placed
  material:           String  (required)             — e.g. "Ceramic", "Stainless Steel"
  colors:             [String]                       — Hex color palette for customization
  customizationAreas: [String]                       — e.g. ["front", "wrap", "handle-side"]
  active:             Boolean (default: true)        — Soft-delete flag
  timestamps:         createdAt, updatedAt           — Auto-managed by Mongoose
}""")
doc.add_paragraph()

h2("7.2 User Schema")
code_block("""\
{
  name:       String   (required)                    — Full name
  email:      String   (required, unique, indexed)   — Login identifier
  password:   String   (required)                    — bcrypt hashed (10 rounds)
  role:       String   (enum: "consumer"|"seller")   — Access level
  phone:      String   (default: "")                 — Optional contact number
  addresses:  [{                                     — Saved delivery addresses
    label:    String,
    line1:    String,
    city:     String,
    state:    String,
    pincode:  String
  }]
  timestamps: createdAt, updatedAt
}""")
doc.add_paragraph()

# ═══════════════════ 8. FEATURES DETAIL ═══════════════════
h1("8. Feature Implementation Details"); add_hr()

h2("8.1 Homepage & Product Pages")
para("The homepage is fully responsive across all screen sizes. On desktop, it features a full-bleed "
     "hero image with overlay text and CTA buttons, a features strip (Premium Packaging, Express "
     "Delivery, Quality Guarantee, Bulk Gifting), a dynamic 'Shop by Gift Type' grid, bestsellers "
     "row ranked by order count, testimonials, and a studio CTA banner. On mobile, a dedicated "
     "MobileHome component renders with touch-optimized horizontal scrollers.")
doc.add_paragraph()

h2("8.2 Authentication Flow")
para("Registration: User submits name, email, password, phone, and role → password hashed with "
     "bcryptjs → user created in MongoDB → JWT generated → HttpOnly cookie set → user object returned.")
doc.add_paragraph()
para("Login: User submits email + password → bcrypt.compare validates → JWT generated → cookie set → "
     "session restored.")
doc.add_paragraph()
para("Session Restore: On every page load, frontend calls GET /api/auth/me → cookie verified → "
     "decoded JWT → user fetched from DB → session state restored without re-login.")
doc.add_paragraph()

h2("8.3 Product Customization Studio")
for b in [
    "Product color selection (swatches from the product's color palette)",
    "Custom text input (up to 72 characters with live preview)",
    "Text color picker (full color wheel)",
    "Placement selector (Top, Center, Bottom)",
    "Font selector (Sans, Serif, Mono, Script)",
    "Quantity input with validation",
    "Image upload (local file via FileReader API with immediate preview)",
]:
    bullet(b)
doc.add_paragraph()

h2("8.4 Performance Optimizations")
for b in [
    "Products pre-loaded with fallback data — Studio renders in <100ms, zero loading time",
    "Session check + product fetch run in parallel via Promise.all",
    "Lite API mode (?lite=true) returns reduced payload for faster network transfer",
    "Unsplash image URLs auto-optimized with width, quality, and format params",
    "ProductCard images use lazy loading with shimmer skeleton placeholder",
]:
    bullet(b)
doc.add_paragraph()

# ═══════════════════ 9. PRODUCT CATALOG ═══════════════════
h1("9. Product Catalog (15 Seeded Products)"); add_hr()
pt = doc.add_table(rows=16, cols=4); pt.style = 'Table Grid'
orange_header(pt, ["Product Name", "Category", "Price", "Material"])
for i, (n,c,p2,m) in enumerate([
    ("Signature Ceramic Mug",      "Mugs",          "Rs. 349",   "Ceramic"),
    ("Steel Insulated Bottle",     "Water Bottles", "Rs. 899",   "Stainless Steel"),
    ("Premium Leather Planner",    "Corporate",     "Rs. 1,299", "Faux Leather"),
    ("Bamboo Desk Organizer",      "Corporate",     "Rs. 1,499", "Bamboo Wood"),
    ("Slim 10k Power Bank",        "Corporate",     "Rs. 1,999", "Aluminum"),
    ("Romantic Anniversary Hamper", "Hampers",      "Rs. 3,999", "Cane Basket"),
    ("3D Photo Moon Lamp",         "Romantic",      "Rs. 2,499", "PLA"),
    ("Ceramic Couple Mug Set",     "Handmade",      "Rs. 1,299", "Stone Clay"),
    ("Luxury Scented Candle Kit",  "Luxury",        "Rs. 1,899", "Soy Wax"),
    ("Pure Silk Sleep Mask",       "Luxury",        "Rs. 999",   "Mulberry Silk"),
    ("Customized Notebook Set",    "Stationery",    "Rs. 799",   "Paper & Board"),
    ("Artisan Jewelry Box",        "Boxes",         "Rs. 1,899", "Wood & Velvet"),
    ("Premium Tea Gift Set",       "Gift Sets",     "Rs. 2,199", "Metal Tins"),
    ("Snap Phone Cover",           "Phone Covers",  "Rs. 449",   "Polycarbonate"),
    ("Gallery Photo Frame",        "Photo Frames",  "Rs. 699",   "Engineered Wood"),
]):
    r = pt.rows[i+1]
    r.cells[0].text = n; r.cells[1].text = c; r.cells[2].text = p2; r.cells[3].text = m
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 10. DEPLOYMENT ═══════════════════
h1("10. Deployment & Hosting"); add_hr()
dt = doc.add_table(rows=6, cols=2); dt.style = 'Table Grid'
orange_header(dt, ["Component", "Details"])
for i, (c2,d) in enumerate([
    ("Live URL",        "https://giftora-six.vercel.app"),
    ("Frontend Host",   "Vercel (automatic deploy on git push to main)"),
    ("Backend / API",   "Vercel Serverless via vercel.json routing"),
    ("Database",        "MongoDB Atlas — cloud cluster, free tier (M0)"),
    ("CI/CD",           "GitHub → Vercel webhook — zero-downtime deployments"),
]):
    r = dt.rows[i+1]
    r.cells[0].text = c2; r.cells[1].text = d
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)
doc.add_paragraph()

# ═══════════════════ 11. CHECKLIST ═══════════════════
h1("11. Milestone 1 Delivery Checklist"); add_hr()
chk = doc.add_table(rows=9, cols=2); chk.style = 'Table Grid'
orange_header(chk, ["Feature", "Status"])
for i, (f,s2) in enumerate([
    ("Homepage and product pages",         "✅ Complete — responsive desktop + mobile"),
    ("User authentication system",         "✅ Complete — JWT, HttpOnly cookies, consumer + seller"),
    ("Product customization module",       "✅ Complete — text, color, font, placement, quantity"),
    ("Image upload functionality",         "✅ Complete — FileReader API with live preview"),
    ("Custom text editing",                "✅ Complete — 72-char limit, 4 fonts, color picker"),
    ("Live product preview system",        "✅ Complete — real-time preview panel"),
    ("Project setup and DB configuration", "✅ Complete — MongoDB Atlas + in-memory fallback"),
    ("Deployment",                         "✅ Live at https://giftora-six.vercel.app"),
]):
    r = chk.rows[i+1]
    r.cells[0].text = f; r.cells[1].text = s2
    r.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0: shade_row(r)

doc.add_paragraph(); doc.add_paragraph()
cl = doc.add_paragraph(); cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cl.add_run("All deliverables for Milestone 1 have been completed and deployed. "
                "The platform is fully functional and ready for Milestone 2.")
cr.italic = True; cr.font.size = Pt(11); cr.font.color.rgb = RGBColor(80,80,80)

doc.save(r"C:\Users\siddh\Desktop\Giftora_Milestone1_Report_v2.docx")
print("Milestone 1 saved!")
